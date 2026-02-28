"""Inspect the sample VAR docx structure."""
import docx
from pathlib import Path

doc = docx.Document('backend/data/sample_var.docx')
lines = []

for ti, t in enumerate(doc.tables):
    lines.append(f'TABLE {ti} ({len(t.rows)}r x {len(t.columns)}c)')
    for row in t.rows:
        lines.append(' | '.join(c.text.strip().replace('\n', ' ')[:40] for c in row.cells))
    lines.append('')

lines.append('=== PARAGRAPHS ===')
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        lines.append(f'{i:3}: [{p.style.name[:18]}] {p.text.strip()[:100]}')

Path('backend/data/var_structure.txt').write_text('\n'.join(lines), encoding='utf-8')
print(f'Done. {len(lines)} lines written to var_structure.txt')
