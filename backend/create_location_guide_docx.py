#!/usr/bin/env python3
"""
Generate SENTRY Location Guide Word Document

Creates a professionally formatted Word document showing all file locations.

Author: Atlas (Code Puppy)
Date: 2026-02-28
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.shared import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from pathlib import Path

# Create document
doc = Document()

# Set default font
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Title
title = doc.add_heading('SENTRY Document Location Guide', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title.runs[0]
title_run.font.color.rgb = RGBColor(0, 83, 226)  # Walmart Blue

# Subtitle
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.add_run('Enterprise Security — Emerging Technology\n').bold = True
subtitle.add_run('Prepared for: Jason Wilbur (j0w16ja)\n')
subtitle.add_run('Date: February 28, 2026\n')
subtitle.add_run('Version: 1.0')

doc.add_paragraph()  # Spacing

# === Quick Reference Map ===
doc.add_heading('🎯 Quick Reference Map', 1)

p = doc.add_paragraph()
p.add_run('PRIMARY LOCATIONS:').bold = True

# Location 1
doc.add_heading('1. SENTRY Application', 2)
p = doc.add_paragraph()
p.add_run('📂 ').bold = True
run = p.add_run('C:\\Users\\j0w16ja\\SENTRY_v2-main')
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0, 83, 226)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Main SENTRY web application (React + FastAPI)')

# Location 2
doc.add_heading('2. VAR Documents (Organized by Month)', 2)
p = doc.add_paragraph()
p.add_run('📂 ').bold = True
run = p.add_run('C:\\Users\\j0w16ja\\OneDrive - Walmart Inc\\ET\\VARs')
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0, 83, 226)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('1,022 Vendor Assessment Reports (202507-202602)')

# Location 3
doc.add_heading('3. Data Files & Trackers', 2)
p = doc.add_paragraph()
p.add_run('📂 ').bold = True
run = p.add_run('C:\\Users\\j0w16ja\\OneDrive - Walmart Inc\\ET\\SENTRY_Data')
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0, 83, 226)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Excel trackers, competitor data, regulatory data, incidents, UAS docs')

# Location 4
doc.add_heading('4. Original Files (Backup)', 2)
p = doc.add_paragraph()
p.add_run('📂 ').bold = True
run = p.add_run('C:\\Users\\j0w16ja\\Downloads')
run.font.name = 'Consolas'
run.font.color.rgb = RGBColor(0, 83, 226)
p = doc.add_paragraph()
p.add_run('Purpose: ').bold = True
p.add_run('Original unorganized files (DO NOT DELETE - safety backup)')

doc.add_page_break()

# === Detailed Folder Structure ===
doc.add_heading('📊 Detailed Folder Structure', 1)

# SENTRY Application
doc.add_heading('1️⃣ SENTRY Application', 2)

table = doc.add_table(rows=1, cols=2)
table.style = 'Light Grid Accent 1'

hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Folder/File'
hdr_cells[1].text = 'Description'

# Make header bold
for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 83, 226)

rows_data = [
    ('backend/data/sentry.db', 'MAIN DATABASE (2,090 vendors)'),
    ('backend/data/sentry.db.backup_20260228', '⭐ SAFETY BACKUP'),
    ('backend/main.py', 'API server (start with: uvicorn main:app)'),
    ('backend/sentry_config.json', '⭐ Document path configuration'),
    ('backend/organize_documents.py', '⭐ Organization script'),
    ('backend/import_enhanced_data_202601_202602.py', '⭐ Data import script'),
    ('backend/health_check.py', '⭐ Verification script'),
    ('components/VendorDashboard.tsx', 'Main vendor directory page'),
    ('components/VendorStatsPanel.tsx', 'KPI dashboard with charts'),
    ('Docs/SENTRY_Organization_Summary.md', '⭐ Full documentation (10 pages)'),
    ('QUICK_REFERENCE.md', '⭐ Quick start guide (1 page)'),
]

for folder, desc in rows_data:
    row_cells = table.add_row().cells
    row_cells[0].text = folder
    row_cells[1].text = desc

p = doc.add_paragraph()
p.add_run('💡 How to Start SENTRY:').bold = True
p = doc.add_paragraph('1. Open terminal in C:\\Users\\j0w16ja\\SENTRY_v2-main', style='List Number')
p = doc.add_paragraph('Run: npm run dev', style='List Number')
p = doc.add_paragraph('Backend auto-starts on port 8000', style='List Number')
p = doc.add_paragraph('Frontend runs on http://localhost:5173', style='List Number')
p = doc.add_paragraph('Click "Vendor Directory" in sidebar', style='List Number')

doc.add_page_break()

# VAR Documents
doc.add_heading('2️⃣ VAR Documents (Organized by Month)', 2)

table2 = doc.add_table(rows=1, cols=3)
table2.style = 'Light Grid Accent 1'

hdr_cells = table2.rows[0].cells
hdr_cells[0].text = 'Month'
hdr_cells[1].text = 'VAR Count'
hdr_cells[2].text = 'Example Files'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 83, 226)

months_data = [
    ('202511/', '80 VARs', 'WMT-SEC-VAR-20251121-DeepKeep-Detailed-v1.docx'),
    ('202512/', '411 VARs ⭐', 'WMT-SEC-VAR-20251208-Axon-Detailed-v1.docx'),
    ('202601/', '366 VARs', 'WMT-SEC-VAR-20260108-DroneShield-Detailed-v1.docx'),
    ('202602/', '159 VARs ⭐ LATEST', 'WMT-SEC-VAR-20260215-Skydio-X10D-Detailed-v1.docx'),
]

for month, count, example in months_data:
    row_cells = table2.add_row().cells
    row_cells[0].text = month
    row_cells[1].text = count
    row_cells[2].text = example

p = doc.add_paragraph()
p.add_run('📊 Total: 1,022 VAR Documents Organized').bold = True

p = doc.add_paragraph()
p.add_run('💡 How to Access:').bold = True
p = doc.add_paragraph('From SENTRY: Click vendor card → "View VAR" button', style='List Bullet')
p = doc.add_paragraph('From File Explorer: Navigate to specific month folder', style='List Bullet')
p = doc.add_paragraph('From SharePoint: Synced via OneDrive', style='List Bullet')

doc.add_page_break()

# Data Files
doc.add_heading('3️⃣ Data Files & Trackers', 2)

# Trackers
doc.add_heading('📁 Trackers/', 3)
p = doc.add_paragraph()
p.add_run('Location: ').bold = True
p.add_run('OneDrive\\ET\\SENTRY_Data\\Trackers\\')
p = doc.add_paragraph()
p.add_run('Contains: ').bold = True
p.add_run('6 Excel files (202507-202602) with vendor data, use cases, maturity levels, ratings')
p = doc.add_paragraph()
p.add_run('Latest: ').bold = True
run = p.add_run('Emerging Tech Tracker_202602.xlsx (260 rows)')
run.font.bold = True

# Competitor Analysis
doc.add_heading('📁 Competitor_Analysis/', 3)
p = doc.add_paragraph()
p.add_run('Location: ').bold = True
p.add_run('OneDrive\\ET\\SENTRY_Data\\Competitor_Analysis\\')
p = doc.add_paragraph()
p.add_run('Contains: ').bold = True
p.add_run('50 CSV files with Amazon, Target, Costco, Kroger events and intel')
p = doc.add_paragraph()
p.add_run('Used by: ').bold = True
p.add_run('CompetitorIntel page in SENTRY')

# Regulatory
doc.add_heading('📁 Regulatory/', 3)
p = doc.add_paragraph()
p.add_run('Location: ').bold = True
p.add_run('OneDrive\\ET\\SENTRY_Data\\Regulatory\\')
p = doc.add_paragraph()
p.add_run('Contains: ').bold = True
p.add_run('30 CSV files with AI laws, biometrics, ALPR, data privacy regulations')
p = doc.add_paragraph()
p.add_run('Used by: ').bold = True
p.add_run('Compliance tracking and risk assessment')

# Incidents
doc.add_heading('📁 Incidents/', 3)
p = doc.add_paragraph()
p.add_run('Location: ').bold = True
p.add_run('OneDrive\\ET\\SENTRY_Data\\Incidents\\')
p = doc.add_paragraph()
p.add_run('Contains: ').bold = True
p.add_run('30 CSV files with ORC, cargo theft, data breaches, violence, cyber incidents')
p = doc.add_paragraph()
p.add_run('Used by: ').bold = True
p.add_run('Incident analysis and threat intelligence')

# UAS Drones
doc.add_heading('📁 UAS_Drones/', 3)
p = doc.add_paragraph()
p.add_run('Location: ').bold = True
p.add_run('OneDrive\\ET\\SENTRY_Data\\UAS_Drones\\')
p = doc.add_paragraph()
p.add_run('Contains: ').bold = True
p.add_run('20 documents (DOCX & PDF) with Skydio, Sunflower Labs, DroneShield, governance docs')
p = doc.add_paragraph()
p.add_run('Used by: ').bold = True
p.add_run('UAS pilot programs and counter-drone assessments')

doc.add_page_break()

# === Navigation Quick Reference ===
doc.add_heading('🗺️ Navigation Quick Reference', 1)

table3 = doc.add_table(rows=1, cols=2)
table3.style = 'Medium Grid 1 Accent 1'

hdr_cells = table3.rows[0].cells
hdr_cells[0].text = 'Task'
hdr_cells[1].text = 'Location'

for cell in hdr_cells:
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
    cell._element.get_or_add_tcPr().append(parse_xml(r'<w:shd {} w:fill="0053E2"/>'.format(nsdecls('w'))))

nav_data = [
    ('Start SENTRY', 'C:\\Users\\j0w16ja\\SENTRY_v2-main\nCommand: npm run dev'),
    ('Find a VAR document', 'OneDrive\\ET\\VARs\\YYYYMM\\\nOr: Click vendor card in SENTRY'),
    ('Get latest vendor data', 'OneDrive\\ET\\SENTRY_Data\\Trackers\\\nFile: Emerging Tech Tracker_202602.xlsx'),
    ('View competitor intel', 'OneDrive\\ET\\SENTRY_Data\\Competitor_Analysis\\'),
    ('Check UAS/Drone docs', 'OneDrive\\ET\\SENTRY_Data\\UAS_Drones\\'),
    ('Import new tracker data', 'SENTRY_v2-main\\backend\\\nCommand: python import_enhanced_...py'),
    ('Organize new downloads', 'SENTRY_v2-main\\backend\\\nCommand: python organize_documents.py'),
    ('Verify data integrity', 'SENTRY_v2-main\\backend\\\nCommand: python health_check.py'),
    ('Restore database backup', 'SENTRY_v2-main\\backend\\data\\\nCopy: sentry.db.backup_20260228'),
]

for task, location in nav_data:
    row_cells = table3.add_row().cells
    row_cells[0].text = task
    row_cells[1].text = location

doc.add_page_break()

# === Monthly Workflow ===
doc.add_heading('🎯 Monthly Workflow', 1)

p = doc.add_paragraph()
p.add_run('When new tracker arrives from SharePoint:').bold = True

doc.add_heading('1️⃣ Download New Tracker', 2)
p = doc.add_paragraph('SharePoint URL: https://teams.wal-mart.com/sites/EmergingTechnologySecurity/Vault/', style='List Bullet')
p = doc.add_paragraph('Download: Emerging Tech Tracker_YYYYMM.xlsx', style='List Bullet')
p = doc.add_paragraph('Save to: C:\\Users\\j0w16ja\\Downloads\\', style='List Bullet')

doc.add_heading('2️⃣ Organize New Files', 2)
p = doc.add_paragraph()
run = p.add_run('cd C:\\Users\\j0w16ja\\SENTRY_v2-main\\backend')
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
run = p.add_run('python organize_documents.py')
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
p.add_run('✅ This copies tracker to: OneDrive\\ET\\SENTRY_Data\\Trackers\\')
p = doc.add_paragraph()
p.add_run('✅ This organizes any new VARs to: OneDrive\\ET\\VARs\\YYYYMM\\')

doc.add_heading('3️⃣ Import Data to Database', 2)
p = doc.add_paragraph()
run = p.add_run('python import_enhanced_data_202601_202602.py')
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
p.add_run('✅ Updates vendor records with latest data')
p = doc.add_paragraph()
p.add_run('✅ Links new VARs to vendor cards')

doc.add_heading('4️⃣ Verify Everything', 2)
p = doc.add_paragraph()
run = p.add_run('python health_check.py')
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
p.add_run('✅ Confirms database integrity')
p = doc.add_paragraph()
p.add_run('✅ Validates document organization')

doc.add_heading('5️⃣ Review in SENTRY', 2)
p = doc.add_paragraph()
run = p.add_run('npm run dev')
run.font.name = 'Consolas'
run.font.size = Pt(10)
p = doc.add_paragraph()
p.add_run('✅ Browse Vendor Directory')
p = doc.add_paragraph()
p.add_run('✅ Verify new data appears')

p = doc.add_paragraph()
p.add_run('⏱️ Total Time: 5 minutes').bold = True

doc.add_page_break()

# === Success Metrics ===
doc.add_heading('📈 Success Metrics', 1)

metrics = [
    '✅ 2,090 Vendors in Database',
    '✅ 1,352 VAR Reports Linked',
    '✅ 45.3% VAR Coverage',
    '✅ 1,022 VAR Documents Organized by Month',
    '✅ 6 Excel Trackers Centralized',
    '✅ 50 Competitor CSVs Ready',
    '✅ 30 Regulatory CSVs Available',
    '✅ 30 Incident CSVs Organized',
    '✅ 20 UAS/Drone Docs Accessible',
    '✅ 46 Technology Categories Tracked',
    '✅ 100% Maturity Data Coverage',
    '✅ Database Integrity 100% Verified',
    '✅ Backup Created & Safe',
]

for metric in metrics:
    p = doc.add_paragraph(metric, style='List Bullet')
    run = p.runs[0]
    run.font.color.rgb = RGBColor(34, 197, 94)  # Green

# === Support ===
doc.add_heading('📞 Support & Contact', 1)

table4 = doc.add_table(rows=5, cols=2)
table4.style = 'Light List Accent 1'

contact_data = [
    ('Owner', 'Jason Wilbur (j0w16ja)'),
    ('Email', 'j0w16ja@walmart.com'),
    ('Team', 'Enterprise Security — Emerging Technology'),
    ('Slack', '#emerging-tech-security'),
    ('Documentation', 'SENTRY_Organization_Summary.md'),
]

for i, (label, value) in enumerate(contact_data):
    row = table4.rows[i]
    row.cells[0].text = label
    row.cells[1].text = value
    row.cells[0].paragraphs[0].runs[0].font.bold = True

# Footer
doc.add_paragraph()
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer.add_run('\n\nDocument Version: 1.0').italic = True
footer.add_run('\nLast Updated: February 28, 2026').italic = True
footer.add_run('\nStatus: ✅ Production Ready').italic = True
footer.add_run('\n\nBuilt with ❤️ by Atlas (Code Puppy) 🐶').italic = True

# Save
output_path = Path(r"C:\Users\j0w16ja\OneDrive - Walmart Inc\ET\SENTRY_Location_Guide.docx")
doc.save(output_path)

print(f"✅ Word document created: {output_path}")
print(f"📄 File size: {output_path.stat().st_size / 1024:.1f} KB")
print("\n🎉 You can now open this document in Microsoft Word!")
