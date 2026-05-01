"""VAR DOCX Score Extractor.

Parses Walmart Vendor Assessment Report DOCX files and extracts:
  - 8 dimension scores (Compliance, Risk, Maturity, Integration, ROI,
    Viability, Differentiation, CloudDep)
  - Composite weighted score
  - Decision band
  - Report date, vendor name

Supports multiple VAR DOCX formats:
  - WMT-SEC-VAR-* format (post-Nov 2025, Code Puppy generated)
  - Vendor-Assessment_* format (earlier Atlas VARs)
  - Vendor Assessment - * format (legacy Word templates)

Usage:
    from var_score_extractor import extract_scores
    scores = extract_scores('path/to/report.docx')
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path
from typing import Optional
from xml.etree import ElementTree as ET

try:
    from docx import Document  # python-docx
except Exception:
    Document = None

# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Canonical dimension name → DB column
DIMENSION_MAP: dict[str, str] = {
    "compliance":      "compliance_score",
    "risk":            "risk_score",
    "maturity":        "maturity_score",
    "integration":     "integration_score",
    "roi":             "roi_score",
    "viability":       "viability_score",
    "differentiation": "differentiation_score",
    "clouddep":        "cloud_dep_score",
    "cloud dep":       "cloud_dep_score",
    "cloud dependency": "cloud_dep_score",
    "cloud dependence": "cloud_dep_score",
}

# Weight mapping for computing composite if table lacks it
WEIGHTS: dict[str, float] = {
    "compliance_score":      0.25,
    "risk_score":            0.25,
    "maturity_score":        0.15,
    "integration_score":     0.10,
    "roi_score":             0.10,
    "viability_score":       0.05,
    "differentiation_score": 0.05,
    "cloud_dep_score":       0.05,
}

SCORE_COL_HEADERS = {
    "score (0", "score (0–5)", "score(0-5)", "score", "dimension score", "raw score"
}

DECISION_BANDS = [
    (4.0, "Advance"),
    (3.0, "Research Further"),
    (2.0, "Defer"),
    (0.0, "Reject"),
]


# ---------------------------------------------------------------------------
# Public API
# ---------------------------------------------------------------------------

def extract_scores(docx_path: str | Path) -> dict:
    """Extract all scoring data from a VAR DOCX.

    Returns a dict with keys matching var_reports table columns:
        overall_score, decision_band,
        compliance_score, risk_score, maturity_score, integration_score,
        roi_score, viability_score, differentiation_score, cloud_dep_score,
        vendor_name, report_date
    Returns empty dict on failure.
    """
    path = Path(docx_path)

    result: dict = {}

    if Document is not None:
        try:
            doc = Document(str(path))
        except Exception:
            doc = None

        if doc is not None:
            result.update(_extract_header(doc))

            table_scores = _extract_from_scoring_table(doc)
            if table_scores:
                result.update(table_scores)
            else:
                result.update(_extract_from_paragraphs(doc))

            if "overall_score" not in result:
                result["overall_score"] = _compute_composite(result)

            if "decision_band" not in result and result.get("overall_score") is not None:
                result["decision_band"] = _score_to_band(result["overall_score"])

            return result

    fallback = _extract_from_docx_xml(path)
    if "_error" in fallback:
        return fallback

    result.update(fallback)

    if "overall_score" not in result:
        result["overall_score"] = _compute_composite(result)
    if "decision_band" not in result and result.get("overall_score") is not None:
        result["decision_band"] = _score_to_band(result["overall_score"])

    return result


# ---------------------------------------------------------------------------
# Internal helpers
# ---------------------------------------------------------------------------

def _extract_header(doc) -> dict:
    """Pull vendor name and report date from opening paragraphs."""
    meta: dict = {}
    for para in doc.paragraphs[:15]:
        text = para.text.strip()
        if not text:
            continue
        # Vendor line: "Vendor: Foo Bar – Product"
        m = re.match(r"Vendor[:：]\s*(.+)", text, re.I)
        if m and "vendor_name" not in meta:
            meta["vendor_name"] = m.group(1).split("–")[0].split("-")[0].strip()

        # Date line: "Date: February 02, 2026" or "2026-02-02"
        m = re.match(r"Date[:：]\s*(.+)", text, re.I)
        if m and "report_date" not in meta:
            meta["report_date"] = m.group(1).strip()
    return meta


def _normalise_dim(name: str) -> Optional[str]:
    """Map a raw dimension name to a DB column name, or None."""
    key = name.lower().strip()
    return DIMENSION_MAP.get(key)


def _parse_score(raw: str) -> Optional[float]:
    """Parse '3.5', '3.5/5', '3.50' etc. → float, or None."""
    raw = raw.strip()
    # Handle fractions like '3.5/5'
    m = re.match(r"([0-9]+\.?[0-9]*)", raw)
    if m:
        val = float(m.group(1))
        if 0.0 <= val <= 5.0:
            return val
    return None


def _extract_from_scoring_table(doc) -> dict:
    """Try to find the dimension-score table and parse it.

    Table 0 in WMT-SEC-VAR format has headers:
      Dimension | Weight (%) | Score (0–5) | Weighted Contribution | ...
    Returns dimension score dict + overall_score + decision_band.
    """
    scores: dict = {}
    weighted_total: Optional[float] = None

    for table in doc.tables:
        if len(table.columns) < 3 or len(table.rows) < 3:
            continue

        # Identify header row
        header = [c.text.strip().lower() for c in table.rows[0].cells]
        if not any(
            any(h in hdr for h in ("score", "weight", "dimension"))
            for hdr in header
        ):
            continue

        # Find which column is the score
        score_col_idx = None
        for ci, hdr in enumerate(header):
            if any(kw in hdr for kw in ("score (0", "score(0", "raw score", "dimension score", "score")):
                score_col_idx = ci
                break
        if score_col_idx is None:
            continue

        # Find weighted contribution column (for sanity check)
        weighted_col_idx = None
        for ci, hdr in enumerate(header):
            if "weighted" in hdr or "contribution" in hdr:
                weighted_col_idx = ci
                break

        dim_col_idx = 0  # Always first column

        dim_scores: dict = {}
        running_weighted = 0.0

        for row in table.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            if not cells:
                continue
            dim_raw = cells[dim_col_idx]
            col = _normalise_dim(dim_raw)
            if col is None:
                continue
            score_raw = cells[score_col_idx] if score_col_idx < len(cells) else ""
            score = _parse_score(score_raw)
            if score is not None:
                dim_scores[col] = score
                # Accumulate weighted total from the contribution column
                if weighted_col_idx is not None and weighted_col_idx < len(cells):
                    wt = _parse_score(cells[weighted_col_idx])
                    if wt is not None:
                        running_weighted += wt

        if len(dim_scores) >= 4:  # Need at least 4 dimensions to trust this table
            scores.update(dim_scores)
            # Use the summed weighted contributions as the overall score
            if running_weighted > 0:
                weighted_total = round(running_weighted, 2)
            break  # Found the scoring table — stop

    if weighted_total is not None:
        scores["overall_score"] = weighted_total
        scores["decision_band"] = _score_to_band(weighted_total)
    elif scores:
        # Compute from dimension scores + known weights
        computed = _compute_composite(scores)
        if computed is not None:
            scores["overall_score"] = computed
            scores["decision_band"] = _score_to_band(computed)

    return scores


def _extract_from_paragraphs(doc) -> dict:
    """Scan paragraphs for 'Composite Weighted Score: X.XX' pattern."""
    scores: dict = {}
    for para in doc.paragraphs:
        text = para.text.strip()
        # "Composite Weighted Score: 3.17 / 5.0 – Decision Band: Research Further"
        m = re.search(
            r"Composite Weighted Score[:：]\s*([0-9]+\.?[0-9]*)\s*/\s*5",
            text, re.I
        )
        if m:
            scores["overall_score"] = round(float(m.group(1)), 2)

        m = re.search(
            r"Decision Band[:：]\s*([\w\s/]+?)(?:\.|$|–|-)",
            text, re.I
        )
        if m and "overall_score" in scores:
            band_raw = m.group(1).strip()
            scores["decision_band"] = _canonicalise_band(band_raw)

    return scores


def _compute_composite(scores: dict) -> Optional[float]:
    """Compute weighted composite from dimension scores."""
    total_weight = 0.0
    total_score = 0.0
    for col, weight in WEIGHTS.items():
        if col in scores:
            total_score += scores[col] * weight
            total_weight += weight
    if total_weight > 0:
        return round(total_score / total_weight * sum(WEIGHTS.values()), 2)
    return None


def _score_to_band(score: float) -> str:
    """Convert numeric score to decision band label."""
    for threshold, band in DECISION_BANDS:
        if score > threshold:
            return band
    return "Reject"


def _canonicalise_band(raw: str) -> str:
    """Normalise messy band strings to canonical form."""
    r = raw.lower().strip()
    if "advance" in r or "pilot" in r:
        return "Advance"
    if "research" in r or "conditional" in r or "further" in r:
        return "Research Further"
    if "defer" in r or "remediat" in r:
        return "Defer"
    if "reject" in r:
        return "Reject"
    return raw.title()


def _extract_from_docx_xml(path: Path) -> dict:
    """Fallback parser for .docx when python-docx is unavailable.

    Reads word/document.xml directly and extracts table-like rows + paragraph text.
    """
    ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    try:
        with zipfile.ZipFile(path, "r") as zf:
            xml_bytes = zf.read("word/document.xml")
    except Exception as exc:  # noqa: BLE001
        return {"_error": f"Unable to read DOCX XML: {exc}"}

    try:
        root = ET.fromstring(xml_bytes)
    except Exception as exc:  # noqa: BLE001
        return {"_error": f"Invalid DOCX XML: {exc}"}

    result: dict = {}

    paragraphs: list[str] = []
    for p in root.findall(".//w:p", ns):
        text = "".join(t.text or "" for t in p.findall(".//w:t", ns)).strip()
        if text:
            paragraphs.append(text)

    for text in paragraphs[:20]:
        vendor_match = re.match(r"Vendor[:：]\s*(.+)", text, re.I)
        if vendor_match and "vendor_name" not in result:
            result["vendor_name"] = vendor_match.group(1).split("–")[0].split("-")[0].strip()

        date_match = re.match(r"Date[:：]\s*(.+)", text, re.I)
        if date_match and "report_date" not in result:
            result["report_date"] = date_match.group(1).strip()

    table_rows: list[list[str]] = []
    for tbl in root.findall(".//w:tbl", ns):
        for tr in tbl.findall(".//w:tr", ns):
            row: list[str] = []
            for tc in tr.findall(".//w:tc", ns):
                txt = "".join(t.text or "" for t in tc.findall(".//w:t", ns)).strip()
                row.append(txt)
            if any(cell for cell in row):
                table_rows.append(row)

    result.update(_extract_from_table_rows(table_rows))
    result.update(_extract_from_text_lines(paragraphs + [" | ".join(r) for r in table_rows]))
    return result


def _extract_from_table_rows(table_rows: list[list[str]]) -> dict:
    scores: dict = {}
    weighted_total: Optional[float] = None

    for idx, row in enumerate(table_rows):
        header = [c.lower() for c in row]
        if not any("dimension" in c for c in header) or not any("score" in c for c in header):
            continue

        score_col_idx = next((i for i, c in enumerate(header) if "score" in c), None)
        dim_col_idx = next((i for i, c in enumerate(header) if "dimension" in c), 0)
        weighted_col_idx = next((i for i, c in enumerate(header) if "weighted" in c or "contribution" in c), None)

        if score_col_idx is None:
            continue

        running_weighted = 0.0
        dim_scores: dict = {}

        for data_row in table_rows[idx + 1 : idx + 20]:
            if max(dim_col_idx, score_col_idx) >= len(data_row):
                continue
            dim_raw = data_row[dim_col_idx]
            col = _normalise_dim(dim_raw)
            if col is None:
                if "decision" in " ".join(data_row).lower() or "composite" in " ".join(data_row).lower():
                    break
                continue

            score = _parse_score(data_row[score_col_idx])
            if score is None:
                continue

            dim_scores[col] = score

            if weighted_col_idx is not None and weighted_col_idx < len(data_row):
                wt = _parse_score(data_row[weighted_col_idx])
                if wt is not None:
                    running_weighted += wt

        if len(dim_scores) >= 4:
            scores.update(dim_scores)
            if running_weighted > 0:
                weighted_total = round(running_weighted, 2)
            break

    if weighted_total is not None:
        scores["overall_score"] = weighted_total
    return scores


def _extract_from_text_lines(lines: list[str]) -> dict:
    scores: dict = {}
    blob = "\n".join(lines)

    overall_match = re.search(
        r"(?:Composite\s+Weighted\s+Score|Overall\s+Score)[:：]?\s*([0-9]+\.?[0-9]*)\s*(?:/\s*5(?:\.0)?)?",
        blob,
        re.I,
    )
    if overall_match:
        scores["overall_score"] = round(float(overall_match.group(1)), 2)

    band_match = re.search(r"Decision\s+Band[:：]?\s*([A-Za-z\-/ ]{3,40})", blob, re.I)
    if band_match:
        raw_band = band_match.group(1).strip()
        norm_band = _canonicalise_band(raw_band)
        if norm_band in {"Advance", "Research Further", "Defer", "Reject"}:
            scores["decision_band"] = norm_band

    for dim_name, col in DIMENSION_MAP.items():
        if col in scores:
            continue
        dim_match = re.search(
            rf"\b{re.escape(dim_name)}\b[^\n\r0-9]{{0,30}}([0-5](?:\.[0-9]+)?)",
            blob,
            re.I,
        )
        if dim_match:
            val = float(dim_match.group(1))
            if 0.0 <= val <= 5.0:
                scores[col] = round(val, 2)

    return scores


# ---------------------------------------------------------------------------
# CLI usage
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import sys, json
    path = sys.argv[1] if len(sys.argv) > 1 else "data/sample_var.docx"
    result = extract_scores(path)
    print(json.dumps(result, indent=2))
